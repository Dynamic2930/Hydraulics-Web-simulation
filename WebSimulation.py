# hydraulic_sim_streamlit.py
import streamlit as st
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import time
import io
from math import pi
import math

st.set_page_config(layout="wide", page_title="Hydraulic Press Simulator")

# -------------------
# Calculation helpers
# -------------------

g = 9.81  # m/s²

# -------------------
# Calculation helpers
# -------------------
def mm_to_m(x):
    return x / 1000.0

def ton_to_kg(t):
    """Convert metric ton → kg (mass)."""
    return t * 1000.0

def area_bore(d_mm):
    """Bore area (m²) from bore diameter in mm."""
    d = mm_to_m(d_mm)
    return (math.pi * (d  ** 2))/4

def area_annulus(bore_mm, rod_mm):
    """Annulus area (m²) from bore & rod diameters in mm."""
    bore = mm_to_m(bore_mm)
    rod = mm_to_m(rod_mm)
    return (math.pi * (bore ** 2 - rod** 2))/4

def pressure_required(force_N, area_m2):

    return (force_N / area_m2)

def pa_to_bar(p):
    return (p * (10**(-5)))

def bar_to_pa(b):
    return (b * (10**(5)))

def flow_rate_Lpm(area_m2, speed_mm_per_s):
    """Flow rate (L/min) from piston area and velocity in mm/s."""
    v = speed_mm_per_s / 1000.0  # → m/s
    q_m3_s = area_m2 * v
    return q_m3_s * 60000.0  # m³/s → L/min

def pump_displacement_cc_per_rev(Q_L_min, rpm, pump_eff):
    """Pump displacement (cc/rev)."""
    return (Q_L_min * 1000.0) / (rpm * pump_eff)

def motor_power_kW(pressure_bar, Q_L_min, pump_eff):
    """Hydraulic motor power in kW."""
    return (((pressure_bar*Q_L_min)/600))

def energy_kJ_from_phase(pressure_bar, Q_L_min, duration_s):
    """Energy consumed in a phase (kJ)."""
    # P = bar_to_pa(pressure_bar)       # Pa
    Q = Q_L_min              # m³/s
    power_W = (pressure_bar * Q  ) /600                # W
    energy_J = power_W * duration_s   # J
    print(power_W,duration_s,energy_J)
    return energy_J         # → kJ


# -------------------
# UI - sidebar inputs
# -------------------
st.title("Hydraulic Press")

with st.sidebar:
    st.header("Cylinder & system inputs")
    bore_mm = st.number_input("Bore (mm)", value=85.0, step=1.0)
    rod_mm = st.number_input("Rod diameter (mm)", value=50.0, step=1.0)
    stroke_mm = st.number_input("Stroke (mm)", value=250.0, step=1.0)
    dead_load_ton = st.number_input("Dead load (ton)", value=1.5, step=0.1)
    holding_load_ton = st.number_input("Holding load (ton)", value=12.0, step=0.1)
    motor_rpm = st.number_input("Motor RPM", value=1800, step=10)
    pump_eff = st.number_input("Pump efficiency", value=0.9, min_value=0.1, max_value=1.0)
    system_loss_bar = st.number_input("System losses (bar)", value=10.0, step=1.0)
    g_val = st.number_input("Gravity g (m/s^2)", value=9.81, format="%.3f")

    st.markdown("---")
    st.header("Duty cycle phases (mm/sec and seconds)")
    # phases: Fast Down, Working, Holding, Fast Up
    fast_down_speed = st.number_input("Fast Down speed (mm/s)", value=200.0)
    fast_down_time = st.number_input("Fast Down time (s)", value=1.0, format="%.2f")
    working_speed = st.number_input("Working speed (mm/s)", value=10.0)
    working_time = st.number_input("Working time (s)", value=5.0, format="%.2f")
    holding_time = st.number_input("Holding time (s)", value=2.0, format="%.2f")
    fast_up_speed = st.number_input("Fast Up speed (mm/s)", value=200.0)
    fast_up_time = st.number_input("Fast Up time (s)", value=1.25, format="%.2f")

    animate_seconds_per_frame = st.slider("Animation: seconds per frame (slower = larger)", 0.02, 1.5, 0.08)

# -------------------
# Calculations
# -------------------
A_bore = area_bore(bore_mm)
A_ann = area_annulus(bore_mm, rod_mm)

dead_force_N = dead_load_ton*1000 * g
holding_force_N = holding_load_ton*1000 * g

working_pressure_pa = pressure_required(holding_force_N, A_bore)  # extension pressure (Pa)
retract_pressure_pa = pressure_required(dead_force_N, A_ann)  # retract pressure using annulus

# Phase-wise pressures (bar)
pressure_fast_down = pa_to_bar(pressure_required(dead_force_N, A_bore)) + system_loss_bar
working_pressure_bar = pa_to_bar(pressure_required(holding_force_N, A_bore)) + system_loss_bar
pressure_holding = working_pressure_bar
retract_pressure_bar = pa_to_bar(pressure_required(dead_force_N, A_ann)) + system_loss_bar



# flows
q_fast_down = flow_rate_Lpm(A_bore, fast_down_speed)
q_working = flow_rate_Lpm(A_bore, working_speed)
q_fast_up = flow_rate_Lpm(A_ann, fast_up_speed)  # retract uses annulus
q_hold = 0.0

max_flow = max(q_fast_down, q_working, q_fast_up)

pump_disp = pump_displacement_cc_per_rev(max_flow, motor_rpm, pump_eff)

motor_kw_fast_down = motor_power_kW(pressure_fast_down, q_fast_down, pump_eff)
motor_kw_working = motor_power_kW(working_pressure_bar, q_working, pump_eff)
motor_kw_holding = motor_power_kW(pressure_holding, q_hold, pump_eff)
motor_kw_fast_up = motor_power_kW(retract_pressure_bar, q_fast_up, pump_eff)

relief_valve_bar = max(working_pressure_bar, retract_pressure_bar) + 20.0

# energy (per phase)
e_fast_down = energy_kJ_from_phase(pressure_fast_down, q_fast_down, fast_down_time)
e_working = energy_kJ_from_phase(working_pressure_bar, q_working, working_time)
e_holding = energy_kJ_from_phase(pressure_holding, q_hold, holding_time)
e_fast_up = energy_kJ_from_phase(retract_pressure_bar, q_fast_up, fast_up_time)

total_energy_kJ = e_fast_down + e_working + e_holding + e_fast_up

# -------------------
# Output panels
# -------------------
col1, col2 = st.columns([1,1])
with col1:
    st.subheader("Cylinder Areas")
    st.write(f"Bore area: **{A_bore:.6f} m²**")
    st.write(f"Annulus area (retract): **{A_ann:.6f} m²**")

    st.subheader("Pressures")
    st.write(f"Pressure (Fast Down): **{pressure_fast_down:.2f} bar**")
    st.write(f"Pressure (Working): **{working_pressure_bar:.2f} bar**")
    st.write(f"Pressure (Holding): **{pressure_holding:.2f} bar**")
    st.write(f"Pressure (Fast Up): **{retract_pressure_bar:.2f} bar**")
    st.write(f"Relief valve suggested: **{relief_valve_bar:.2f} bar**")

with col2:
    st.subheader("Flow & Sizing")
    st.write(f"Fast down flow: **{q_fast_down:.2f} L/min**")
    st.write(f"Working flow: **{q_working:.2f} L/min**")
    st.write(f"Fast up (retract) flow: **{q_fast_up:.2f} L/min**")
    st.write(f"Max flow (for pump sizing): **{max_flow:.2f} L/min**")
    st.write(f"Pump displacement (cc/rev): **{pump_disp:.2f} cc/rev**")
    # st.write(f"Estimated motor power: **{motor_kw:.2f} kW**")

   
    st.write(f"Motor power (Fast Down): **{motor_kw_fast_down:.2f} kW**")
    st.write(f"Motor power (Working): **{motor_kw_working:.2f} kW**")
    st.write(f"Motor power (Holding): **{motor_kw_holding:.2f} kW**")
    st.write(f"Motor power (Fast Up): **{motor_kw_fast_up:.2f} kW**")


st.markdown("---")
st.subheader("Energy per phase (kJ)")
st.write(pd.DataFrame({
    "Phase":["Fast Down","Working","Holding","Fast Up","Total"],
    "Duration_s":[fast_down_time, working_time, holding_time, fast_up_time, fast_down_time+working_time+holding_time+fast_up_time],
    "Energy_kJ":[e_fast_down, e_working, e_holding, e_fast_up, total_energy_kJ]
}).set_index("Phase"))

# -------------------
# Build time-series for a single cycle
# -------------------
def build_cycle_ts():
    # we will construct arrays with small dt
    dt = 0.05  # resolution in seconds
    phases = [
        ("Fast Down", fast_down_time, -fast_down_speed, A_bore, pressure_fast_down, q_fast_down),
        ("Working", working_time, -working_speed, A_bore, working_pressure_bar, q_working),
        ("Holding", holding_time, 0.0, A_bore, pressure_holding, q_hold),
        ("Fast Up", fast_up_time, fast_up_speed, A_ann, retract_pressure_bar, q_fast_up)
    ]
    times = []
    strokes = []
    flows = []
    pressures = []
    t = 0.0
    stroke_position_mm = stroke_mm  # start at top (idle); we will simulate a cycle: down => up
    direction = -1.0
    for name, duration, spd, area_use, pres_bar, q in phases:
        n_steps = max(1, int(duration / dt))
        for i in range(n_steps):
            # update stroke
            stroke_position_mm += spd * dt
            # clamp 0..stroke
            stroke_position_mm = max(0.0, min(stroke_mm, stroke_position_mm))
            t += dt
            times.append(t)
            strokes.append(stroke_position_mm)
            flows.append(q)
            pressures.append(pres_bar)
    return np.array(times), np.array(strokes), np.array(flows), np.array(pressures)

times, strokes, flows, pressures = build_cycle_ts()

st.subheader("Cycle plot (animated)")
placeholder = st.empty()


fig, ax = plt.subplots(3,1, figsize=(8,8), sharex=True)
ax[0].plot(times, strokes, color="tab:blue")
ax[0].set_ylabel("Stroke (mm)")
ax[0].invert_yaxis()
ax[0].grid(True, linestyle="--", alpha=0.6)

ax[1].plot(times, flows, color="tab:green")
ax[1].set_ylabel("Flow (L/min)")
ax[1].grid(True, linestyle="--", alpha=0.6)

ax[2].plot(times, pressures, color="tab:red")
ax[2].set_ylabel("Pressure (bar)")
ax[2].set_xlabel("Time (s)")
ax[2].grid(True, linestyle="--", alpha=0.6)

plt.tight_layout(pad=2.0)
fig_path = "cycle_plot.png"
fig.savefig(fig_path)
plt.close(fig)


# Animation controls with phase shading
if st.button("Run slow animation"):
    anim_placeholder = st.empty()
    frames = len(times)

    # Define colors for curves
    stroke_color = "tab:blue"
    flow_color   = "tab:green"
    press_color  = "tab:red"

    # Phase timing boundaries
    t_fast_down_end = fast_down_time
    t_working_end   = t_fast_down_end + working_time
    t_holding_end   = t_working_end + holding_time
    t_fast_up_end   = t_holding_end + fast_up_time

    for i in range(0,frames,2):
        fig2, a2 = plt.subplots(3, 1, figsize=(9, 9), sharex=True)

        # Common phase shading
        for ax in a2:
            ax.axvspan(0, t_fast_down_end, color="orange", alpha=0.15, label="Fast Down" if i == 0 else "")
            ax.axvspan(t_fast_down_end, t_working_end, color="red", alpha=0.15, label="Working" if i == 0 else "")
            ax.axvspan(t_working_end, t_holding_end, color="purple", alpha=0.15, label="Holding" if i == 0 else "")
            ax.axvspan(t_holding_end, t_fast_up_end, color="green", alpha=0.15, label="Fast Up" if i == 0 else "")

        # --- Stroke Plot ---
        a2[0].plot(times[:i+1], strokes[:i+1], color=stroke_color, linewidth=2)
        a2[0].scatter(times[i], strokes[i], color=stroke_color, s=50, zorder=5)
        a2[0].set_ylabel("Stroke (mm)")
        a2[0].set_title("Cylinder Stroke vs Time")
        a2[0].invert_yaxis()
        a2[0].grid(True, linestyle="--", alpha=0.6)

        # --- Flow Plot ---
        a2[1].plot(times[:i+1], flows[:i+1], color=flow_color, linewidth=2)
        a2[1].scatter(times[i], flows[i], color=flow_color, s=50, zorder=5)
        a2[1].set_ylabel("Flow (L/min)")
        a2[1].set_title("Pump Flow vs Time")
        a2[1].grid(True, linestyle="--", alpha=0.6)

        # --- Pressure Plot ---
        a2[2].plot(times[:i+1], pressures[:i+1], color=press_color, linewidth=2)
        a2[2].scatter(times[i], pressures[i], color=press_color, s=50, zorder=5)
        a2[2].set_ylabel("Pressure (bar)")
        a2[2].set_title("System Pressure vs Time")
        a2[2].set_xlabel("Time (s)")
        a2[2].grid(True, linestyle="--", alpha=0.6)

        # Only add phase legend once
        if i == 0:
            a2[0].legend(loc="upper right")

        plt.tight_layout(pad=2.0)
        anim_placeholder.pyplot(fig2)
        plt.close(fig2)

        # time.sleep(0.1)

    st.success("✅ Animation finished with phase highlights!")

# -------------------
# Export: CSV
# -------------------
st.markdown("---")
st.subheader("Export data")
df_cycle = pd.DataFrame({"time_s": times, "stroke_mm": strokes, "flow_L_min": flows, "pressure_bar": pressures})
buf = io.StringIO()
df_cycle.to_csv(buf, index=False)
st.download_button("Download cycle CSV", data=buf.getvalue(), file_name="hydraulic_cycle.csv", mime="text/csv")

st.info("Use the inputs to tweak parameters and re-run. Animation uses a simple frame loop for smooth playback.")

from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from reportlab.lib.styles import ParagraphStyle
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
from reportlab.lib.units import cm


# Register Times New Roman
pdfmetrics.registerFont(TTFont('TimesNewRoman', 'times.ttf'))  # requires times.ttf in your system or project folder

def export_pdf(inputs, results, df_energy, df_phase, fig_path):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)

    # Custom styles
    base_style = ParagraphStyle(
        "Base",
        fontName="TimesNewRoman",
        fontSize=18,
        leading=22,   # line spacing
    )
    heading_style = ParagraphStyle(
        "Heading",
        parent=base_style,
        fontName="TimesNewRoman",
        fontSize=20,
        leading=24,
        spaceAfter=12,
    )
    title_style = ParagraphStyle(
        "Title",
        parent=base_style,
        fontName="TimesNewRoman",
        fontSize=24,
        leading=28,
        alignment=1,  # center
        spaceAfter=20,
    )

    story = []

    # --- Title ---
    story.append(Paragraph("Hydraulic Press Simulation Report", title_style))
    story.append(Spacer(1, 12))

    # --- Inputs ---
    story.append(Paragraph("Inputs", heading_style))
    for k, v in inputs.items():
        story.append(Paragraph(f"{k}: {v}", base_style))
    story.append(Spacer(1, 12))

    # --- Results ---
    story.append(Paragraph("Results", heading_style))
    for k, v in results.items():
        story.append(Paragraph(f"{k}: {v}", base_style))
    story.append(Spacer(1, 12))

    # --- Formulas ---
    story.append(Paragraph("Formulas Used", heading_style))
    formulas = [
        "Cylinder Bore Area:  A_bore = π × (D_bore²) / 4",
        "Annulus Area:  A_ann = π × (D_bore² – D_rod²) / 4",
        "Pressure:  P = F / A",
        "Flow Rate:  Q = A × v   (converted to L/min)",
        "Pump Displacement:  Disp = (Q × 1000) / (RPM × η)",
        "Motor Power:  P_kW = (p_bar × Q_L/min) / 600",
        "Energy:  E_kJ = (p × Q × t) / 1000"
    ]
    for f in formulas:
        story.append(Paragraph(f, base_style))
    story.append(Spacer(1, 12))

 
# --- Phase Summary Table ---
    story.append(Paragraph("Phase Summary", heading_style))

    # Convert DataFrame to list of lists
    data = [df_phase_out.columns.to_list()] + df_phase_out.values.tolist()

    # Set fixed column widths (adjust as needed)
    col_widths = [3*cm, 3*cm, 3*cm, 4*cm, 3*cm]

    tbl = Table(data, colWidths=col_widths, hAlign="LEFT")

    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0,0), (-1,0), colors.grey),
        ("TEXTCOLOR", (0,0), (-1,0), colors.whitesmoke),
        ("ALIGN", (0,0), (-1,-1), "CENTER"),
        ("FONTNAME", (0,0), (-1,-1), "Times-Roman"),
        ("FONTSIZE", (0,0), (-1,0), 14),   # Header row
        ("FONTSIZE", (0,1), (-1,-1), 12), # Table body
        ("BOTTOMPADDING", (0,0), (-1,0), 8),
        ("GRID", (0,0), (-1,-1), 0.5, colors.black),
    ]))
    story.append(tbl)
    story.append(Spacer(1, 12))
    # --- Graph ---
    story.append(Paragraph("Cycle Graphs", heading_style))
    story.append(RLImage(fig_path, width=400, height=300))

    doc.build(story)
    buffer.seek(0)
    return buffer

# # -------------------
# # Export Word
# # -------------------
# def export_word(inputs, results, df_energy, fig_path):
#     doc = Document()
#     doc.add_heading("Hydraulic Press Simulation Report", 0)

#     doc.add_heading("Inputs", level=1)
#     for k, v in inputs.items():
#         doc.add_paragraph(f"{k}: {v}")

#     doc.add_heading("Results", level=1)
#     for k, v in results.items():
#         doc.add_paragraph(f"{k}: {v}")

#     doc.add_heading("Energy per Phase", level=1)
#     table = doc.add_table(rows=1, cols=len(df_energy.columns))
#     hdr_cells = table.rows[0].cells
#     for i, col in enumerate(df_energy.columns):
#         hdr_cells[i].text = str(col)
#     for row in df_energy.values.tolist():
#         row_cells = table.add_row().cells
#         for i, val in enumerate(row):
#             row_cells[i].text = str(val)

#     doc.add_heading("Cycle Graphs", level=1)
#     doc.add_picture(fig_path, width=docx.shared.Inches(5))

#     buffer = io.BytesIO()
#     doc.save(buffer)
#     buffer.seek(0)
#     return buffer

# -------------------
# Buttons in Streamlit
# -------------------
st.markdown("---")
st.subheader("Export PDF / Word")

# Save last plot as image
fig_path = "cycle_plot.png"
fig.savefig(fig_path)

inputs_dict = {
    "Bore (mm)": bore_mm,
    "Rod (mm)": rod_mm,
    "Stroke (mm)": stroke_mm,
    "Dead load (ton)": dead_load_ton,
    "Holding load (ton)": holding_load_ton,
    "Motor RPM": motor_rpm,
    "Pump efficiency": pump_eff,
    "System loss (bar)": system_loss_bar,
    "Gravity g (m/s²)": g_val
}
results_dict = {
    "Pressure Fast Down (bar)": f"{pressure_fast_down:.2f}",
    "Motor Power Fast Down (kW)": f"{motor_kw_fast_down:.2f}",
    "Pressure Working (bar)": f"{working_pressure_bar:.2f}",
    "Motor Power Working (kW)": f"{motor_kw_working:.2f}",
    "Pressure Holding (bar)": f"{pressure_holding:.2f}",
    "Motor Power Holding (kW)": f"{motor_kw_holding:.2f}",
    "Pressure Fast Up (bar)": f"{retract_pressure_bar:.2f}",
    "Motor Power Fast Up (kW)": f"{motor_kw_fast_up:.2f}",
    "Pump displacement (cc/rev)": f"{pump_disp:.2f}",
    "Relief valve (bar)": f"{relief_valve_bar:.2f}"
}

df_energy_out = pd.DataFrame({
    "Duration (s)":[fast_down_time, working_time, holding_time, fast_up_time, fast_down_time+working_time+holding_time+fast_up_time],
    "Energy (kJ)":[e_fast_down, e_working, e_holding, e_fast_up, total_energy_kJ]
}, index=["Fast Down","Working","Holding","Fast Up","Total"])

df_phase_out = pd.DataFrame({
    "Phase": ["Fast Down", "Working", "Holding", "Fast Up"],
    "Time (s)": [fast_down_time, working_time, holding_time, fast_up_time],
    "Pressure (bar)": [pressure_fast_down, working_pressure_bar, pressure_holding, retract_pressure_bar],
    "Motor Power (kW)": [motor_kw_fast_down, motor_kw_working, motor_kw_holding, motor_kw_fast_up],
    "Energy (kJ)": [e_fast_down, e_working, e_holding, e_fast_up]
})
df_phase_out = df_phase_out.round(2)
pdf_buf = export_pdf(inputs_dict, results_dict, df_energy_out, df_phase_out, fig_path)

st.download_button("Download PDF Report", data=pdf_buf, file_name="hydraulic_report.pdf", mime="application/pdf")

# word_buf = export_word(inputs_dict, results_dict, df_energy_out, fig_path)
# st.download_button("Download Word Report", data=word_buf, file_name="hydraulic_report.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# -------------------
# Import CSV
# -------------------
st.markdown("---")
st.subheader("Import Cycle CSV")
uploaded_csv = st.file_uploader("Upload CSV", type=["csv"])
if uploaded_csv:
    df_uploaded = pd.read_csv(uploaded_csv)
    st.write("Imported data:", df_uploaded.head())
    st.line_chart(df_uploaded.set_index("time_s"))